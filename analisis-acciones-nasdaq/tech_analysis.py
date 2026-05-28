# -*- coding: utf-8 -*-
"""Download real data for 5 tech stocks, run a transparent fundamental
scoring, and emit a DOCX report + a JSON with the scored data.
Outputs are written next to this script, so it runs on any machine."""
import os
import json
import math
import datetime as dt
import numpy as np
import yfinance as yf

TICKERS = ["NVDA", "MSFT", "AAPL", "GOOGL", "META"]
HERE = os.path.dirname(os.path.abspath(__file__))
JSON_OUT = os.path.join(HERE, "data.json")
DOCX_OUT = os.path.join(HERE, "Analisis_Tecnologicas.docx")


def g(info, *keys):
    for k in keys:
        v = info.get(k)
        if v is not None:
            return v
    return None


def fetch(ticker):
    t = yf.Ticker(ticker)
    info = t.info
    hist = t.history(period="1y")
    closes = hist["Close"].dropna()
    rets = closes.pct_change().dropna()
    year_ret = float(closes.iloc[-1] / closes.iloc[0] - 1) if len(closes) > 1 else None
    vol = float(rets.std() * math.sqrt(252)) if len(rets) > 2 else None
    sma50 = float(closes.tail(50).mean()) if len(closes) >= 50 else None
    sma200 = float(closes.tail(200).mean()) if len(closes) >= 200 else None
    price = g(info, "currentPrice", "regularMarketPrice") or (
        float(closes.iloc[-1]) if len(closes) else None
    )
    target = g(info, "targetMeanPrice")
    upside = (target / price - 1) if (target and price) else None
    return {
        "ticker": ticker,
        "name": g(info, "shortName", "longName") or ticker,
        "price": price,
        "marketCap": g(info, "marketCap"),
        "trailingPE": g(info, "trailingPE"),
        "forwardPE": g(info, "forwardPE", "trailingPE"),
        "priceToBook": g(info, "priceToBook"),
        "roe": g(info, "returnOnEquity"),
        "margin": g(info, "profitMargins"),
        "revGrowth": g(info, "revenueGrowth"),
        "earnGrowth": g(info, "earningsGrowth"),
        "dividendYield": g(info, "dividendYield"),
        "beta": g(info, "beta"),
        "debtToEquity": g(info, "debtToEquity"),
        "target": target,
        "upside": upside,
        "yearReturn": year_ret,
        "volatility": vol,
        "sma50": sma50,
        "sma200": sma200,
        "recKey": g(info, "recommendationKey"),
    }


def norm(values, higher_better=True):
    """Min-max normalize a list (None -> neutral 0.5)."""
    nums = [v for v in values if v is not None]
    if not nums:
        return [0.5 for _ in values]
    lo, hi = min(nums), max(nums)
    out = []
    for v in values:
        if v is None or hi == lo:
            out.append(0.5)
        else:
            s = (v - lo) / (hi - lo)
            out.append(s if higher_better else 1 - s)
    return out


def score(stocks):
    # factor -> (key, higher_better, weight)
    factors = [
        ("valuation", "forwardPE", False, 0.20),
        ("growth_rev", "revGrowth", True, 0.13),
        ("growth_earn", "earnGrowth", True, 0.12),
        ("roe", "roe", True, 0.14),
        ("margin", "margin", True, 0.11),
        ("health", "debtToEquity", False, 0.15),
        ("upside", "upside", True, 0.15),
    ]
    comp = {s["ticker"]: 0.0 for s in stocks}
    breakdown = {s["ticker"]: {} for s in stocks}
    for name, key, hb, w in factors:
        vals = [s[key] for s in stocks]
        ns = norm(vals, hb)
        for s, n in zip(stocks, ns):
            comp[s["ticker"]] += n * w
            breakdown[s["ticker"]][name] = round(n * 100, 1)
    for s in stocks:
        s["score"] = round(comp[s["ticker"]] * 100, 1)
        s["breakdown"] = breakdown[s["ticker"]]
    stocks.sort(key=lambda s: s["score"], reverse=True)
    for i, s in enumerate(stocks):
        s["rank"] = i + 1
    return stocks


def notes(s):
    """Data-driven qualitative bullets."""
    out = []
    if s["forwardPE"]:
        if s["forwardPE"] < 22:
            out.append(f"Valuación razonable (P/U fwd {s['forwardPE']:.1f}).")
        elif s["forwardPE"] > 35:
            out.append(f"Valuación exigente (P/U fwd {s['forwardPE']:.1f}): el mercado descuenta crecimiento.")
        else:
            out.append(f"Valuación media (P/U fwd {s['forwardPE']:.1f}).")
    if s["roe"]:
        out.append(f"Rentabilidad {'muy alta' if s['roe']>0.3 else 'sólida' if s['roe']>0.15 else 'moderada'} (ROE {s['roe']*100:.0f}%).")
    if s["revGrowth"] is not None:
        out.append(f"Crecimiento de ingresos {s['revGrowth']*100:+.0f}% interanual.")
    if s["debtToEquity"] is not None:
        out.append(f"Apalancamiento {'bajo' if s['debtToEquity']<60 else 'moderado' if s['debtToEquity']<120 else 'alto'} (deuda/capital {s['debtToEquity']:.0f}).")
    if s["upside"] is not None:
        out.append(f"Potencial vs. precio objetivo de analistas: {s['upside']*100:+.0f}%.")
    if s["yearReturn"] is not None:
        out.append(f"Rendimiento 12 meses: {s['yearReturn']*100:+.0f}%.")
    return out


def fmt_cap(v):
    if not v:
        return "n/d"
    return f"${v/1e12:.2f} B" if v >= 1e12 else f"${v/1e9:.0f} mil M"


def pct(v, plus=False):
    if v is None:
        return "n/d"
    return f"{v*100:+.1f}%" if plus else f"{v*100:.1f}%"


def num(v, d=1):
    return "n/d" if v is None else f"{v:.{d}f}"


def build_docx(stocks, asof):
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    h = doc.add_heading("Análisis fundamental: 5 acciones tecnológicas", level=0)
    sub = doc.add_paragraph(f"Datos al {asof} · Fuente: Yahoo Finance (yfinance)")
    sub.runs[0].italic = True

    doc.add_paragraph(
        "Este informe descarga datos de mercado y razones financieras de cinco "
        "tecnológicas, las puntúa con un modelo transparente y concluye con la "
        "acción mejor posicionada según fundamentales."
    )

    doc.add_heading("Metodología", level=1)
    doc.add_paragraph(
        "Cada acción recibe una puntuación compuesta (0-100) combinando 7 factores, "
        "normalizados entre las 5 empresas: valuación (P/U adelantado, 20%), "
        "crecimiento de ingresos (13%) y de utilidades (12%), rentabilidad sobre "
        "capital o ROE (14%), margen de ganancia (11%), salud financiera o "
        "deuda/capital (15%) y potencial frente al precio objetivo de analistas (15%)."
    )

    # Comparison table
    doc.add_heading("Tabla comparativa", level=1)
    cols = ["Acción", "Precio", "Cap.", "P/U fwd", "ROE", "Margen", "Crec. ing.", "Deuda/Cap", "Potencial", "Puntaje"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].paragraphs[0].add_run(c).bold = True
    for s in stocks:
        row = table.add_row().cells
        vals = [
            s["ticker"],
            f"${s['price']:.0f}" if s["price"] else "n/d",
            fmt_cap(s["marketCap"]),
            num(s["forwardPE"]),
            pct(s["roe"]),
            pct(s["margin"]),
            pct(s["revGrowth"], True),
            num(s["debtToEquity"], 0),
            pct(s["upside"], True),
            f"{s['score']:.0f}",
        ]
        for i, v in enumerate(vals):
            row[i].paragraphs[0].add_run(str(v))

    # Per-stock detail
    doc.add_heading("Detalle por acción", level=1)
    for s in stocks:
        doc.add_heading(f"#{s['rank']} · {s['name']} ({s['ticker']}) — puntaje {s['score']:.0f}/100", level=2)
        for n in notes(s):
            doc.add_paragraph(n, style="List Bullet")

    # Conclusion
    win = stocks[0]
    doc.add_heading("Conclusión y recomendación", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Acción recomendada para comprar: {win['name']} ({win['ticker']}).").bold = True
    doc.add_paragraph(
        f"{win['ticker']} obtiene el puntaje más alto ({win['score']:.0f}/100) al combinar "
        f"de forma equilibrada rentabilidad, crecimiento, valuación y potencial al alza. "
        f"Le siguen {stocks[1]['ticker']} ({stocks[1]['score']:.0f}) y {stocks[2]['ticker']} ({stocks[2]['score']:.0f})."
    )

    doc.add_heading("Aviso", level=1)
    d = doc.add_paragraph(
        "Este documento es material educativo y de análisis. NO constituye asesoría "
        "financiera ni una recomendación personalizada de compra o venta. Los datos "
        "provienen de Yahoo Finance y pueden contener errores o estar desactualizados; "
        "el rendimiento pasado no garantiza resultados futuros. Toma tus decisiones de "
        "inversión con tu propio criterio y, si lo necesitas, con un asesor certificado."
    )
    d.runs[0].italic = True

    os.makedirs(os.path.dirname(DOCX_OUT), exist_ok=True)
    doc.save(DOCX_OUT)


def main():
    asof = dt.date.today().strftime("%d/%m/%Y")
    stocks = []
    for tk in TICKERS:
        print("fetching", tk)
        stocks.append(fetch(tk))
    stocks = score(stocks)

    os.makedirs(os.path.dirname(JSON_OUT), exist_ok=True)
    with open(JSON_OUT, "w") as f:
        json.dump({"asof": asof, "stocks": stocks}, f, ensure_ascii=False, indent=2)

    build_docx(stocks, asof)

    print("\n=== RANKING ===")
    for s in stocks:
        print(f"#{s['rank']} {s['ticker']:5} score={s['score']:.0f}  PEfwd={num(s['forwardPE'])}  ROE={pct(s['roe'])}  revG={pct(s['revGrowth'],True)}  up={pct(s['upside'],True)}")
    print("\nGanadora:", stocks[0]["ticker"])
    print("DOCX:", DOCX_OUT)
    print("JSON:", JSON_OUT)


if __name__ == "__main__":
    main()
